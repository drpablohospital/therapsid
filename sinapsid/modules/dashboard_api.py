#!/usr/bin/env python3
"""
Dashboard UCI Completo - Métricas de primer mundo
Basado en NHSN, Leapfrog, AHRQ y SCM de España
"""

from flask import Blueprint, jsonify, request, render_template
from datetime import datetime, timedelta
from collections import Counter, defaultdict
import statistics
from modules.database import get_db_connection
from modules.clinical_analytics import calculate_advanced_metrics
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

dashboard_bp = Blueprint('dashboard', __name__, url_prefix='/uci-dashboard')

DEFAULT_TOTAL_BEDS = 20


def get_patients_filtered(start_date=None, end_date=None, servicio=None, procedencia=None):
    """Obtener pacientes con filtros opcionales"""
    with get_db_connection() as conn:
        cur = conn.cursor()
        
        query = """SELECT * FROM patients WHERE 1=1"""
        params = []
        
        if start_date and end_date:
            query += """ AND (fecha_ingreso BETWEEN %s AND %s 
                OR (fecha_egreso_uci IS NULL OR fecha_egreso_uci BETWEEN %s AND %s))"""
            params.extend([start_date, end_date, start_date, end_date])
        
        if servicio:
            query += " AND servicio_tratante = %s"
            params.append(servicio)
        
        if procedencia:
            query += " AND procedencia = %s"
            params.append(procedencia)
        
        cur.execute(query, params)
        columns = [desc[0] for desc in cur.description]
        patients = [dict(zip(columns, row)) for row in cur.fetchall()]
        cur.close()
    return patients


def get_evolutions_filtered(patient_ids=None, start_date=None, end_date=None):
    """Obtener evoluciones filtradas"""
    with get_db_connection() as conn:
        cur = conn.cursor()
        query = """SELECT * FROM evoluciones WHERE 1=1"""
        params = []
        
        if patient_ids:
            placeholders = ','.join(['%s'] * len(patient_ids))
            query += f" AND patient_id IN ({placeholders})"
            params.extend(patient_ids)
        
        if start_date and end_date:
            query += " AND fecha BETWEEN %s AND %s"
            params.extend([start_date, end_date])
        
        query += " ORDER BY fecha, hora"
        cur.execute(query, params)
        columns = [desc[0] for desc in cur.description]
        evolutions = [dict(zip(columns, row)) for row in cur.fetchall()]
        cur.close()
    return evolutions


def get_cultivos_data(start_date=None, end_date=None):
    """Obtener datos de cultivos"""
    with get_db_connection() as conn:
        cur = conn.cursor()
        if start_date and end_date:
            cur.execute("SELECT * FROM cultivos WHERE fecha BETWEEN %s AND %s", (start_date, end_date))
        else:
            cur.execute("SELECT * FROM cultivos")
        columns = [desc[0] for desc in cur.description]
        cultivos = [dict(zip(columns, row)) for row in cur.fetchall()]
        cur.close()
    return cultivos


def calculate_device_days(patients, evolutions, device_field):
    """Calcular días-dispositivo para tasas de infección"""
    total_days = 0
    for p in patients:
        patient_evols = [e for e in evolutions if e.get('patient_id') == p.get('id')]
        for e in patient_evols:
            if e.get(device_field):
                total_days += 1
    return total_days


def calculate_global_metrics(patients, evolutions, cultivos, total_beds=DEFAULT_TOTAL_BEDS):
    """Calcular métricas globales del dashboard"""
    
    total_patients = len(patients)
    if total_patients == 0:
        return {}
    
    # ===== MÉTRICAS DE POBLACIÓN =====
    active_patients = [p for p in patients if p.get('estado') == 'ingreso']
    discharged_patients = [p for p in patients if p.get('estado') == 'egreso']
    
    # Mortalidad
    deceased_patients = [p for p in patients if 'fallecido' in str(p.get('condicion_egreso', '')).lower() 
                         or 'defuncion' in str(p.get('condicion_egreso', '')).lower()]
    mortality_rate = (len(deceased_patients) / total_patients * 100) if total_patients > 0 else 0
    
    # ===== MÉTRICAS DE ESTANCIA =====
    dias_estancia_list = [p.get('dias_estancia', 0) or 0 for p in patients]
    avg_dias_estancia = sum(dias_estancia_list) / total_patients if total_patients > 0 else 0
    median_dias_estancia = statistics.median(dias_estancia_list) if dias_estancia_list else 0
    
    # Pacientes > 7 días (ALOS prolongado)
    long_stay_patients = [d for d in dias_estancia_list if d > 7]
    long_stay_rate = (len(long_stay_patients) / total_patients * 100) if total_patients > 0 else 0
    
    # Pacientes > 14 días (estancia muy prolongada)
    very_long_stay = [d for d in dias_estancia_list if d > 14]
    very_long_stay_rate = (len(very_long_stay) / total_patients * 100) if total_patients > 0 else 0
    
    # ===== MÉTRICAS DE OCUPACIÓN =====
    occupied_beds = len(active_patients)
    occupancy_rate = (occupied_beds / total_beds * 100) if total_beds > 0 else 0
    available_beds = total_beds - occupied_beds
    
    # ===== MÉTRICAS CLÍNICAS AVANZADAS =====
    total_with_vm = 0
    total_with_shock = 0
    total_with_sepsis = 0
    total_with_sdr = 0
    total_with_lra = 0
    total_with_coma = 0
    total_with_coagulopathy = 0
    total_with_hyperlactatemia = 0
    
    # Métricas de severidad promedio
    sofa_scores = []
    apache_scores = []
    glasgow_scores = []
    
    # Dispositivos
    total_cvc_days = 0
    total_ett_days = 0
    total_urinary_days = 0
    total_arterial_days = 0
    total_hemodialysis_days = 0
    
    patient_ids = [p.get('id') for p in patients]
    patient_evolutions = {pid: [] for pid in patient_ids}
    for e in evolutions:
        pid = e.get('patient_id')
        if pid in patient_evolutions:
            patient_evolutions[pid].append(e)
    
    for patient in patients:
        pid = patient.get('id')
        patient_evol = patient_evolutions.get(pid, [])
        
        # Scores de severidad
        if patient.get('sofa2_ingreso'):
            sofa_scores.append(patient.get('sofa2_ingreso'))
        if patient.get('apache2_ingreso'):
            apache_scores.append(patient.get('apache2_ingreso'))
        if patient.get('glasgow'):
            glasgow_scores.append(patient.get('glasgow'))
        
        # Calcular métricas avanzadas
        if patient_evol:
            try:
                metrics = calculate_advanced_metrics(patient, patient_evol)
                
                if metrics.get('vm_required'):
                    total_with_vm += 1
                if metrics.get('shock_detected'):
                    total_with_shock += 1
                if metrics.get('sepsis_detected'):
                    total_with_sepsis += 1
                if metrics.get('sdr_detected'):
                    total_with_sdr += 1
                if metrics.get('lra_detected'):
                    total_with_lra += 1
                if metrics.get('coma_detected'):
                    total_with_coma += 1
                if metrics.get('coagulopathy'):
                    total_with_coagulopathy += 1
                if metrics.get('hyperlactatemia'):
                    total_with_hyperlactatemia += 1
            except Exception:
                pass
        
        # Días de dispositivos (usar campos de evoluciones)
        for e in patient_evol:
            if e.get('cateter_cvc'):
                total_cvc_days += 1
            if e.get('tubo_endotraqueal'):
                total_ett_days += 1
            if e.get('sonda_urinaria'):
                total_urinary_days += 1
            if e.get('linea_arterial'):
                total_arterial_days += 1
            if e.get('cateter_hemodialisis'):
                total_hemodialysis_days += 1
    
    # Promedios de scores
    avg_sofa = sum(sofa_scores) / len(sofa_scores) if sofa_scores else 0
    avg_apache = sum(apache_scores) / len(apache_scores) if apache_scores else 0
    avg_glasgow = sum(glasgow_scores) / len(glasgow_scores) if glasgow_scores else 0
    
    # ===== MÉTRICAS DE CULTIVOS E INFECCIONES =====
    positive_cultivos = [c for c in cultivos if c.get('resultado') == 'Positivo']
    negative_cultivos = [c for c in cultivos if c.get('resultado') == 'Negativo']
    pending_cultivos = [c for c in cultivos if c.get('resultado') == 'Pendiente']
    
    # Tasas de infección por 1000 días-dispositivo (NHSN style)
    clabsi_rate = (len([c for c in positive_cultivos if 'sangre' in str(c.get('tipo', '')).lower()]) / total_cvc_days * 1000) if total_cvc_days > 0 else 0
    cauti_rate = (len([c for c in positive_cultivos if 'orina' in str(c.get('tipo', '')).lower()]) / total_urinary_days * 1000) if total_urinary_days > 0 else 0
    vapi_rate = (len([c for c in positive_cultivos if 'bronquial' in str(c.get('tipo', '')).lower() or 'tráquea' in str(c.get('tipo', '')).lower()]) / total_ett_days * 1000) if total_ett_days > 0 else 0
    
    # Microorganismos
    microorganismos = Counter()
    for c in positive_cultivos:
        micro = c.get('microorganismo')
        if micro and micro not in [None, 'None', '']:
            microorganismos[micro.upper()] += 1
    top_microorganismos = microorganismos.most_common(10)
    
    # Resistencias
    resistencias = Counter()
    for c in positive_cultivos:
        res = c.get('resistencia')
        if res and res not in [None, 'None', '']:
            antibiotics = [r.strip() for r in str(res).split(',')]
            for antibiotic in antibiotics:
                if antibiotic and antibiotic not in ['None', '']:
                    resistencias[antibiotic] += 1
    top_resistencias = resistencias.most_common(10)
    
    # Tipos de cultivos
    tipos_cultivo = Counter()
    for c in cultivos:
        tipo = c.get('tipo')
        if tipo:
            tipos_cultivo[tipo] += 1
    
    # ===== MÉTRICAS DEMOGRÁFICAS =====
    edades = [p.get('edad', 0) or 0 for p in patients]
    avg_edad = sum(edades) / len(edades) if edades else 0
    
    sexos = Counter([p.get('sexo') for p in patients if p.get('sexo')])
    
    # ===== MÉTRICAS DE READMISIÓN =====
    readmissions = len([p for p in patients if (p.get('episodio') and p.get('episodio').startswith('Reingreso')) or 'reingreso' in str(p.get('padecimiento_actual', '')).lower()])
    readmission_rate = (readmissions / total_patients * 100) if total_patients > 0 else 0
    
    return {
        # === MÉTRICAS DE POBLACIÓN ===
        'total_patients': total_patients,
        'active_patients': len(active_patients),
        'discharged_patients': len(discharged_patients),
        'deceased_patients': len(deceased_patients),
        'mortality_rate': round(mortality_rate, 1),
        'readmission_rate': round(readmission_rate, 1),
        'readmissions': readmissions,
        
        # === MÉTRICAS DE ESTANCIA (ALOS) ===
        'avg_dias_estancia': round(avg_dias_estancia, 1),
        'median_dias_estancia': round(median_dias_estancia, 1),
        'long_stay_rate': round(long_stay_rate, 1),
        'long_stay_count': len(long_stay_patients),
        'very_long_stay_rate': round(very_long_stay_rate, 1),
        'very_long_stay_count': len(very_long_stay),
        
        # === MÉTRICAS DE OCUPACIÓN ===
        'occupancy_rate': round(occupancy_rate, 1),
        'total_beds': total_beds,
        'occupied_beds': occupied_beds,
        'available_beds': available_beds,
        
        # === MÉTRICAS CLÍNICAS ===
        'vm_rate': round(total_with_vm / total_patients * 100, 1) if total_patients > 0 else 0,
        'vm_count': total_with_vm,
        'shock_rate': round(total_with_shock / total_patients * 100, 1) if total_patients > 0 else 0,
        'shock_count': total_with_shock,
        'sepsis_rate': round(total_with_sepsis / total_patients * 100, 1) if total_patients > 0 else 0,
        'sepsis_count': total_with_sepsis,
        'sdr_rate': round(total_with_sdr / total_patients * 100, 1) if total_patients > 0 else 0,
        'sdr_count': total_with_sdr,
        'lra_rate': round(total_with_lra / total_patients * 100, 1) if total_patients > 0 else 0,
        'lra_count': total_with_lra,
        'coma_rate': round(total_with_coma / total_patients * 100, 1) if total_patients > 0 else 0,
        'coma_count': total_with_coma,
        'coagulopathy_rate': round(total_with_coagulopathy / total_patients * 100, 1) if total_patients > 0 else 0,
        'coagulopathy_count': total_with_coagulopathy,
        'hyperlactatemia_rate': round(total_with_hyperlactatemia / total_patients * 100, 1) if total_patients > 0 else 0,
        'hyperlactatemia_count': total_with_hyperlactatemia,
        
        # === SCORES DE SEVERIDAD ===
        'avg_sofa': round(avg_sofa, 1),
        'avg_apache': round(avg_apache, 1),
        'avg_glasgow': round(avg_glasgow, 1),
        
        # === MÉTRICAS DE DISPOSITIVOS ===
        'total_cvc_days': total_cvc_days,
        'total_ett_days': total_ett_days,
        'total_urinary_days': total_urinary_days,
        'total_arterial_days': total_arterial_days,
        'total_hemodialysis_days': total_hemodialysis_days,
        
        # === MÉTRICAS DE INFECCIÓN (NHSN) ===
        'clabsi_rate': round(clabsi_rate, 2),
        'cauti_rate': round(cauti_rate, 2),
        'vapi_rate': round(vapi_rate, 2),
        'total_cultivos': len(cultivos),
        'positive_cultivos': len(positive_cultivos),
        'negative_cultivos': len(negative_cultivos),
        'pending_cultivos': len(pending_cultivos),
        'positive_cultivo_rate': round(len(positive_cultivos) / len(cultivos) * 100, 1) if len(cultivos) > 0 else 0,
        
        # === MICROORGANISMOS Y RESISTENCIAS ===
        'top_microorganismos': [
            {'name': micro, 'count': count, 'percentage': round(count / len(positive_cultivos) * 100, 1)}
            for micro, count in top_microorganismos
        ],
        'top_resistencias': [
            {'antibiotic': res, 'count': count, 'percentage': round(count / len(positive_cultivos) * 100, 1)}
            for res, count in top_resistencias
        ],
        'tipos_cultivo': dict(tipos_cultivo),
        
        # === DEMOGRAFÍA ===
        'avg_edad': round(avg_edad, 1),
        'sex_distribution': dict(sexos),
    }


@dashboard_bp.route('/data')
def dashboard_data():
    """Endpoint principal del dashboard"""
    
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    servicio = request.args.get('servicio')
    procedencia = request.args.get('procedencia')
    total_beds = request.args.get('total_beds', type=int) or DEFAULT_TOTAL_BEDS
    
    # Obtener datos
    patients = get_patients_filtered(start_date, end_date, servicio, procedencia)
    patient_ids = [p.get('id') for p in patients]
    evolutions = get_evolutions_filtered(patient_ids, start_date, end_date)
    cultivos = get_cultivos_data(start_date, end_date)
    
    # Calcular métricas
    metrics = calculate_global_metrics(patients, evolutions, cultivos, total_beds)
    
    # Agregar filtros aplicados
    metrics['filters'] = {
        'start_date': start_date,
        'end_date': end_date,
        'servicio': servicio,
        'procedencia': procedencia,
        'total_beds': total_beds
    }
    
    return jsonify(metrics)


@dashboard_bp.route('/')
def dashboard_view():
    """Renderizar la vista del dashboard"""
    return render_template('dashboard_integrated.html')

import csv
import io
from flask import Response

# Endpoint para exportar datos a CSV
@dashboard_bp.route('/export')
def export_data():
    """Exportar TODOS los datos en formato tidy data (estándar de investigación)."""
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    servicio = request.args.get('servicio')
    procedencia = request.args.get('procedencia')
    
    # Obtener pacientes según filtros
    with get_db_connection() as conn:
        cur = conn.cursor()
        
        query = "SELECT id FROM patients WHERE 1=1"
        params = []
        
        if start_date:
            query += " AND fecha_ingreso >= %s"
            params.append(start_date)
        if end_date:
            query += " AND fecha_ingreso <= %s"
            params.append(end_date)
        if servicio:
            query += " AND servicio_tratante = %s"
            params.append(servicio)
        if procedencia:
            query += " AND procedencia = %s"
            params.append(procedencia)
        
        cur.execute(query, tuple(params))
        patient_ids = [str(row[0]) for row in cur.fetchall()]
        
        if not patient_ids:
            cur.close()
            return 'No hay datos para exportar', 404
        
        placeholders = ','.join(['%s'] * len(patient_ids))
        
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Encabezados estándar tidy data
        writer.writerow(['patient_id', 'tipo_registro', 'fecha', 'hora', 'variable', 'valor', 'unidad', 'tabla_origen', 'notas'])
        
        # 1. PATIENTS - Datos demográficos y de ingreso
        cur.execute(f"SELECT * FROM patients WHERE id IN ({placeholders})", tuple(patient_ids))
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        
        for row in rows:
            patient_id = row[columns.index('id')]
            fecha_ingreso = row[columns.index('fecha_ingreso')] if 'fecha_ingreso' in columns else None
            
            for i, col in enumerate(columns):
                if col in ['id', 'created_at', 'updated_at']:
                    continue
                valor = row[i]
                if valor is not None:
                    writer.writerow([
                        patient_id, 'paciente', fecha_ingreso, '', col, str(valor), '', 'patients', ''
                    ])
        
        # 2. EVOLUCIONES
        cur.execute(f"SELECT * FROM evoluciones WHERE patient_id IN ({placeholders}) ORDER BY patient_id, fecha DESC", tuple(patient_ids))
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        
        for row in rows:
            patient_id = row[columns.index('patient_id')]
            fecha = row[columns.index('fecha')]
            hora = row[columns.index('hora')] if 'hora' in columns else ''
            
            for i, col in enumerate(columns):
                if col in ['id', 'patient_id']:
                    continue
                valor = row[i]
                if valor is not None:
                    writer.writerow([
                        patient_id, 'evolucion', fecha, hora, col, str(valor), '', 'evoluciones', ''
                    ])
        
        # 3. CULTIVOS
        cur.execute(f"SELECT * FROM cultivos WHERE patient_id IN ({placeholders}) ORDER BY patient_id, fecha DESC", tuple(patient_ids))
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        
        for row in rows:
            patient_id = row[columns.index('patient_id')]
            fecha = row[columns.index('fecha')]
            
            for i, col in enumerate(columns):
                if col in ['id', 'patient_id']:
                    continue
                valor = row[i]
                if valor is not None:
                    writer.writerow([
                        patient_id, 'cultivo', fecha, '', col, str(valor), '', 'cultivos', ''
                    ])
        
        # 4. CLINICAL ANALYTICS
        cur.execute(f"SELECT * FROM clinical_analytics WHERE patient_id IN ({placeholders})", tuple(patient_ids))
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        
        for row in rows:
            patient_id = row[columns.index('patient_id')]
            created_at = row[columns.index('created_at')]
            
            for i, col in enumerate(columns):
                if col in ['id', 'patient_id', 'created_at', 'updated_at']:
                    continue
                valor = row[i]
                if valor is not None:
                    writer.writerow([
                        patient_id, 'analytics',
                        created_at.date() if created_at else '',
                        created_at.strftime('%H:%M:%S') if created_at else '',
                        col, str(valor), '', 'clinical_analytics', ''
                    ])
        
        # 5. CLINICAL SNAPSHOTS
        cur.execute(f"SELECT * FROM clinical_snapshots WHERE patient_id IN ({placeholders}) ORDER BY patient_id, snapshot_date DESC", tuple(patient_ids))
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        
        for row in rows:
            patient_id = row[columns.index('patient_id')]
            snapshot_date = row[columns.index('snapshot_date')]
            
            for i, col in enumerate(columns):
                if col in ['id', 'patient_id']:
                    continue
                valor = row[i]
                if valor is not None:
                    writer.writerow([
                        patient_id, 'snapshot',
                        snapshot_date.date() if snapshot_date else '',
                        snapshot_date.strftime('%H:%M:%S') if snapshot_date else '',
                        col, str(valor), '', 'clinical_snapshots', ''
                    ])
        
        # 6. CLINICAL NOTES
        cur.execute(f"SELECT * FROM clinical_notes WHERE patient_id IN ({placeholders}) ORDER BY patient_id, created_at DESC", tuple(patient_ids))
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        
        for row in rows:
            patient_id = row[columns.index('patient_id')]
            created_at = row[columns.index('created_at')]
            
            for i, col in enumerate(columns):
                if col in ['id', 'patient_id', 'created_at']:
                    continue
                valor = row[i]
                if valor is not None:
                    writer.writerow([
                        patient_id, 'nota_clinica',
                        created_at.date() if created_at else '',
                        created_at.strftime('%H:%M:%S') if created_at else '',
                        col, str(valor), '', 'clinical_notes', ''
                    ])
        
        # 7. EVOLUCION TEXTO LIBRE
        cur.execute(f"SELECT * FROM evolucion_texto_libre WHERE patient_id IN ({placeholders}) ORDER BY patient_id, created_at DESC", tuple(patient_ids))
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        
        for row in rows:
            patient_id = row[columns.index('patient_id')]
            created_at = row[columns.index('created_at')]
            
            for i, col in enumerate(columns):
                if col in ['id', 'patient_id', 'created_at']:
                    continue
                valor = row[i]
                if valor is not None:
                    writer.writerow([
                        patient_id, 'evolucion_texto',
                        created_at.date() if created_at else '',
                        created_at.strftime('%H:%M:%S') if created_at else '',
                        col, str(valor), '', 'evolucion_texto_libre', ''
                    ])
        
        # 8. MEDICAMENTOS
        med_tables = {
            'medicamentos_gastro': 'med_gastro',
            'medicamentos_hemodinamicos': 'med_hemodinamico',
            'medicamentos_nefro': 'med_nefro',
            'medicamentos_neurologicos': 'med_neurologico',
            'medicacion_hematologica': 'med_hematologico'
        }
        
        for table, tipo_reg in med_tables.items():
            cur.execute(f"SELECT * FROM {table} WHERE patient_id IN ({placeholders})", tuple(patient_ids))
            columns = [desc[0] for desc in cur.description]
            rows = cur.fetchall()
            
            for row in rows:
                patient_id = row[columns.index('patient_id')]
                
                for i, col in enumerate(columns):
                    if col in ['id', 'patient_id']:
                        continue
                    valor = row[i]
                    if valor is not None:
                        writer.writerow([
                            patient_id, tipo_reg, '', '', col, str(valor), '', table, ''
                        ])
        
        # 9. ANTIBIOTICOS
        cur.execute(f"SELECT * FROM antibioticos WHERE patient_id IN ({placeholders})", tuple(patient_ids))
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        
        for row in rows:
            patient_id = row[columns.index('patient_id')]
            
            for i, col in enumerate(columns):
                if col in ['id', 'patient_id']:
                    continue
                valor = row[i]
                if valor is not None:
                    writer.writerow([
                        patient_id, 'antibiotico', '', '', col, str(valor), '', 'antibioticos', ''
                    ])
        
        # 10. TRANSFUSIONES
        cur.execute(f"SELECT * FROM transfusiones WHERE patient_id IN ({placeholders})", tuple(patient_ids))
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        
        for row in rows:
            patient_id = row[columns.index('patient_id')]
            
            for i, col in enumerate(columns):
                if col in ['id', 'patient_id']:
                    continue
                valor = row[i]
                if valor is not None:
                    writer.writerow([
                        patient_id, 'transfusion', '', '', col, str(valor), '', 'transfusiones', ''
                    ])
        
        cur.close()
    
    csv_data = output.getvalue()
    output.close()
    
    return Response(csv_data, mimetype='text/csv', headers={
        'Content-Disposition': f'attachment; filename=sinapsid_uci_tidy_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
    })

def generate_census_doc():
    """Generar censo institucional UCI con formato idéntico al documento original."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    with get_db_connection() as conn:
        cur = conn.cursor()
        cur.execute("""
            SELECT id, nombre_completo, edad, curp, fecha_nacimiento, expediente,
                   diagnostico_ingreso, fecha_ingreso_hosp, fecha_ingreso, cama, estado,
                   modo_ventilatorio, fio2, peep, vt_ml, ps_cmh2o
            FROM patients 
            WHERE estado = 'ingreso'
            ORDER BY cama, id
        """)
        patients = cur.fetchall()
        
        if not patients:
            return None
        
        # Crear documento con orientación horizontal (Landscape)
        doc = Document()
        section = doc.sections[0]
        section.orientation = 1  # LANDSCAPE
        section.page_width = Inches(11.0)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.433)
        section.right_margin = Inches(0.433)
        section.top_margin = Inches(0.433)
        section.bottom_margin = Inches(0.433)
        
        # Título: CENSO DE PACIENTES EN UCI
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('CENSO DE PACIENTES EN UCI')
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Fecha
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(datetime.now().strftime('%d/%m/%Y'))
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()  # Espacio
        
        # Crear tabla con 11 columnas
        headers = [
            'CAMA', 'NOMBRE, EDAD, CURP, F. NACIMIENTO, EXPEDIENTE', 'DIAGNÓSTICO',
            'FI: HOSP', 'FI: UCI', 'ESTADO DE SALUD', 'PRONÓSTICO', 'AMV',
            'ANTIBIÓTICOS / CULTIVOS / PATOLOGÍA', 'PENDIENTES', 'LABORATORIOS RELEVANTES'
        ]
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configurar anchos de columna exactos (en pulgadas)
        col_widths = [
            0.394, 1.476, 2.091, 0.642, 0.646,
            0.594, 0.973, 0.590, 1.181, 0.788, 1.083
        ]
        
        # Encabezados
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            # Configurar ancho
            hdr_cells[i].width = Inches(col_widths[i])
            
            # Texto del encabezado
            paragraph = hdr_cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(header)
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Fondo gris para encabezados
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D9D9D9')
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            
            # Bordes
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)
        
        # Llenar datos de cada paciente
        for patient in patients:
            (patient_id, nombre, edad, curp, fecha_nac, expediente,
             diagnostico_ingreso, fecha_ingreso_hosp, fecha_ingreso, cama, estado,
             modo_vent, fio2, peep, vt_ml, ps_cmh2o) = patient
            
            # Obtener última evolución para labs
            cur.execute("""
                SELECT hemoglobina, leucocitos, plaquetas, glucosa, bun, creatinina,
                       sodio, potasio, cloro, calcio, magnesio, fosforo
                FROM evoluciones 
                WHERE patient_id = %s 
                ORDER BY fecha DESC, hora DESC 
                LIMIT 1
            """, (patient_id,))
            last_evo = cur.fetchone()
            
            # Obtener scores
            cur.execute("""
                SELECT sofa_baseline, apache2_baseline
                FROM clinical_analytics 
                WHERE patient_id = %s
                ORDER BY created_at DESC 
                LIMIT 1
            """, (patient_id,))
            scores = cur.fetchone()
            
            # Obtener cultivos recientes
            cur.execute("""
                SELECT tipo, resultado, fecha, microorganismo
                FROM cultivos 
                WHERE patient_id = %s 
                ORDER BY fecha DESC 
                LIMIT 3
            """, (patient_id,))
            cultivos = cur.fetchall()
            
            # Obtener antibioticos activos
            cur.execute("""
                SELECT antibiotico, via, dosis, unidad
                FROM antibioticos 
                WHERE patient_id = %s
            """, (patient_id,))
            antibioticos = cur.fetchall()
            
            # Crear fila
            row_cells = table.add_row().cells
            
            # Configurar anchos
            for i, width in enumerate(col_widths):
                row_cells[i].width = Inches(width)
            
            # 1. CAMA (7pt)
            paragraph = row_cells[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(cama) if cama else '')
            run.font.name = 'Calibri'
            run.font.size = Pt(7)
            
            # 2. NOMBRE, EDAD, CURP, F. NACIMIENTO, EXPEDIENTE (7pt)
            paragraph = row_cells[1].paragraphs[0]
            fecha_nac_str = fecha_nac.strftime('%d/%m/%Y') if fecha_nac else ''
            run = paragraph.add_run(f"{nombre}\n{edad} AÑOS\nCURP: {curp or ''}\nFN: {fecha_nac_str}\nEXP: {expediente or ''}")
            run.font.name = 'Calibri'
            run.font.size = Pt(7)
            
            # 3. DIAGNÓSTICO (6pt)
            paragraph = row_cells[2].paragraphs[0]
            if diagnostico_ingreso:
                # Separar diagnósticos por líneas
                diag_lines = diagnostico_ingreso.split('\n')
                for j, line in enumerate(diag_lines):
                    run = paragraph.add_run(line.strip())
                    run.font.name = 'Calibri'
                    run.font.size = Pt(6)
                    if j < len(diag_lines) - 1:
                        run = paragraph.add_run()
            else:
                run = paragraph.add_run('')
                run.font.name = 'Calibri'
                run.font.size = Pt(6)
            
            # 4. FI: HOSP (6pt)
            paragraph = row_cells[3].paragraphs[0]
            fecha_hosp_str = fecha_ingreso_hosp.strftime('%d/%m/%y') if fecha_ingreso_hosp else ''
            run = paragraph.add_run(fecha_hosp_str)
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            
            # 5. FI: UCI (6pt)
            paragraph = row_cells[4].paragraphs[0]
            fecha_uci_str = fecha_ingreso.strftime('%d/%m/%y') if fecha_ingreso else ''
            run = paragraph.add_run(fecha_uci_str)
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            
            # 6. ESTADO DE SALUD (6pt, default "Delicado")
            paragraph = row_cells[5].paragraphs[0]
            run = paragraph.add_run("Delicado")
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            
            # 7. PRONÓSTICO (6pt)
            paragraph = row_cells[6].paragraphs[0]
            pronostico = ""
            if scores:
                sofa, apache = scores
                if sofa is not None:
                    pronostico += f"SOFA: {sofa}\n"
                if apache is not None:
                    # Calcular mortalidad aproximada según APACHE II
                    if apache <= 10:
                        mort = "< 10%"
                    elif apache <= 20:
                        mort = "10-40%"
                    elif apache <= 30:
                        mort = "40-70%"
                    else:
                        mort = "> 70%"
                    pronostico += f"APACHE II: {apache} puntos, Mortalidad: {mort}\n"
            run = paragraph.add_run(pronostico.strip() if pronostico else "Sin datos")
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            
            # 8. AMV (6pt)
            paragraph = row_cells[7].paragraphs[0]
            amv = ""
            if modo_vent:
                amv += f"{modo_vent}\n"
            if fio2:
                amv += f"FiO2: {fio2}%\n"
            if peep:
                amv += f"PEEP: {peep}\n"
            if vt_ml:
                amv += f"VT: {vt_ml}\n"
            if ps_cmh2o:
                amv += f"PS: {ps_cmh2o}\n"
            run = paragraph.add_run(amv.strip() if amv else "No")
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            
            # 9. ANTIBIÓTICOS / CULTIVOS / PATOLOGÍA (6pt)
            paragraph = row_cells[8].paragraphs[0]
            cultivo_text = ""
            
            # Cultivos
            if cultivos:
                for tipo, resultado, fecha, microorganismo in cultivos:
                    fecha_str = fecha.strftime('%d/%m/%y') if fecha else ''
                    if microorganismo:
                        cultivo_text += f"{tipo} {fecha_str}: {microorganismo}\n"
                    else:
                        cultivo_text += f"{tipo} {fecha_str}: {resultado}\n"
            
            # Antibióticos
            if antibioticos:
                cultivo_text += "\nANTIBIÓTICOS:\n"
                for antibiotico, via, dosis, unidad in antibioticos:
                    cultivo_text += f"- {antibiotico} ({via}) {dosis} {unidad}\n"
            
            run = paragraph.add_run(cultivo_text.strip() if cultivo_text else "Sin datos")
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            
            # 10. PENDIENTES (vacío, 6pt)
            paragraph = row_cells[9].paragraphs[0]
            run = paragraph.add_run("")
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            
            # 11. LABORATORIOS RELEVANTES (6pt)
            paragraph = row_cells[10].paragraphs[0]
            labs = ""
            if last_evo:
                (hemoglobina, leucocitos, plaquetas, glucosa, bun, creatinina,
                 sodio, potasio, cloro, calcio, magnesio, fosforo) = last_evo
                
                if hemoglobina:
                    labs += f"HB {hemoglobina} "
                if leucocitos:
                    labs += f"LEU {leucocitos} "
                if plaquetas:
                    labs += f"PLT {plaquetas}\n"
                if glucosa:
                    labs += f"GLU {glucosa} "
                if bun:
                    labs += f"BUN {bun} "
                if creatinina:
                    labs += f"CRT {creatinina}\n"
                if sodio:
                    labs += f"NA {sodio} "
                if potasio:
                    labs += f"K {potasio} "
                if cloro:
                    labs += f"CL {cloro}\n"
                if calcio:
                    labs += f"CA {calcio} "
                if magnesio:
                    labs += f"MG {magnesio} "
                if fosforo:
                    labs += f"P {fosforo}"
            
            run = paragraph.add_run(labs.strip() if labs else "Sin labs recientes")
            run.font.name = 'Calibri'
            run.font.size = Pt(6)
            
            # Aplicar bordes a todas las celdas de la fila
            for cell in row_cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
        
        cur.close()
    
    # Guardar documento en memoria
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io
@dashboard_bp.route('/censo')
def export_censo():
    """Exportar censo institucional de pacientes activos."""
    doc = generate_census_doc()
    if not doc:
        return 'No hay pacientes con estado de ingreso', 404
    
    return Response(
        doc.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={
            'Content-Disposition': f'attachment; filename=censo_uci_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        }
    )

